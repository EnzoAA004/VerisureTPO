from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"C:\Users\enzoa\Downloads\TP_Integrador_Verisure_Smart_Twin_10_hojas.docx")


def add_p(doc: Document, text: str = ""):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(3)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(11, 19, 36)
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(229, 0, 0)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("UADE\nUniversidad Argentina de la Empresa\n")
    run.bold = True
    run.font.size = Pt(13)
    run = cover.add_run("Trabajo Practico Integrador\n")
    run.bold = True
    run.font.size = Pt(16)
    run = cover.add_run("Verisure Smart-Twin\n")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(229, 0, 0)
    run = cover.add_run("Orquestador multimodal predictivo para seguridad inteligente")
    run.italic = True
    run.font.size = Pt(11)

    add_table(
        doc,
        ["Campo", "Descripcion"],
        [
            ("Materia", "Tendencias Tecnologicas"),
            ("Empresa base", "Verisure Argentina"),
            ("Proyecto", "Verisure Smart-Twin"),
            ("Tipo", "Propuesta academica de mejora tecnologica + MVP funcional"),
            ("Integrantes", "[Completar]"),
            ("Fecha", "[Completar]"),
            ("Aclaracion", "Propuesta academica no oficial. No representa una implementacion real ni autorizada por Verisure."),
        ],
        [4, 12],
    )

    doc.add_heading("Resumen ejecutivo", level=1)
    add_p(doc, "El trabajo propone Verisure Smart-Twin, una mejora academica aplicada a Verisure Argentina para asistir a una Central Receptora de Alarmas (CRA). El problema identificado es el alto volumen de eventos recibidos, incluyendo falsas alarmas o situaciones ambiguas que consumen tiempo operativo y pueden demorar la atencion de intrusiones reales.")
    add_p(doc, "La solucion se plantea como un orquestador multimodal predictivo que combina senales visuales, auditivas, sensoriales, conectividad y contexto. Con esos datos calcula un Risk Score entre 0 y 100, clasifica cada evento como verde, amarillo o rojo, recomienda una accion para el operador y alimenta un tablero gerencial de Business Intelligence.")
    add_p(doc, "El proyecto integra contenidos centrales de la materia: TIC/NTIC, Inteligencia Artificial, Internet de las Cosas, Cloud Computing, Business Intelligence, BYOD, Realidad Aumentada, ciberseguridad y privacidad. El MVP local desarrollado en React permite simular eventos, observar la decision del sistema y ver indicadores gerenciales alimentados por las simulaciones de la sesion.")

    doc.add_heading("1. Empresa, identidad y definicion estrategica", level=1)
    add_p(doc, "La empresa base seleccionada es Verisure Argentina, del sector terciario de servicios tecnologicos de seguridad. Su rubro principal es la seguridad electronica, monitoreo de alarmas, domotica aplicada, dispositivos conectados y atencion remota de eventos.")
    add_p(doc, "Para el trabajo se utiliza una identidad academica propia: Verisure Smart-Twin. No se usan logos oficiales ni recursos marcarios descargados. La paleta visual se inspira en rojo de alerta, azul noche tecnologico y gris plata. Slogan propuesto: “Seguridad inteligente para decidir antes, responder mejor y proteger mas”.")
    add_table(
        doc,
        ["Elemento", "Definicion sintetica"],
        [
            ("Mision", "Proteger hogares, comercios y organizaciones mediante monitoreo inteligente, respuesta profesional y uso responsable de tecnologia."),
            ("Vision", "Evolucionar desde un modelo reactivo de alarma hacia una seguridad predictiva, trazable y basada en datos."),
            ("Objetivos", "Reducir falsas alarmas, priorizar eventos criticos, mejorar tiempos de respuesta, generar indicadores BI y abrir planes Premium/Agro."),
            ("Valores", "Proteccion, innovacion responsable, privacidad, transparencia, eficiencia y trabajo en equipo."),
        ],
        [4, 12],
    )

    doc.add_heading("2. Productos, servicios y mejora propuesta", level=1)
    add_p(doc, "Verisure opera como un ecosistema de seguridad monitoreada: sensores, paneles, camaras/fotodetectores, conectividad, app movil, central de monitoreo, soporte tecnico y protocolos de respuesta. Smart-Twin no reemplaza estos servicios, sino que agrega una capa de inteligencia operativa para interpretar mejor cada evento.")
    add_table(
        doc,
        ["Servicio actual", "Funcion", "Mejora Smart-Twin"],
        [
            ("CRA 24/7", "Recepcion y gestion de alarmas.", "Priorizacion por Risk Score, semaforo de estado y accion recomendada."),
            ("Sensores IoT", "Movimiento, apertura, shock, inhibicion o perimetro.", "Cruce con audio, visual, horario, zona y conectividad."),
            ("Camaras/fotodetectores", "Evidencia visual.", "Clasificacion simulada de riesgo visual sin reconocimiento facial."),
            ("ZeroVision", "Disuasion por humo antirrobo.", "Hand-off sensorial cuando la vision queda bloqueada."),
            ("App movil", "Control remoto y notificaciones.", "Historial inteligente, explicaciones y recomendaciones futuras."),
            ("Planes Premium/Agro", "Mayor cobertura o valor agregado.", "Reportes, LoRaWAN, satelital y sensores perimetrales."),
        ],
        [3.2, 4.8, 7.8],
    )

    doc.add_heading("3. Mercado objetivo", level=1)
    add_p(doc, "El mercado natural se concentra en grandes centros urbanos como CABA, Gran Buenos Aires, Cordoba y Rosario, donde existen hogares, comercios y PyMEs con capacidad de abono mensual y alta sensibilidad frente a la inseguridad. La propuesta tambien contempla expansion agro/rural, donde la conectividad es limitada, las distancias son mayores y la cobertura perimetral es critica.")
    add_table(
        doc,
        ["Segmento", "Necesidad", "Propuesta de valor"],
        [
            ("Hogar urbano base", "Reducir falsas alarmas domesticas y proteger la vivienda.", "Clasificacion de eventos simples y trazabilidad basica."),
            ("Hogar Premium", "Mayor inteligencia, evidencia y respuesta.", "Analisis multimodal, hand-off sensorial y reportes."),
            ("PyME/comercio", "Proteccion de local, mercaderia y empleados.", "Priorizacion, auditoria y tablero operativo."),
            ("Agro/rural", "Cobertura de grandes extensiones y baja conectividad.", "LoRaWAN, satelital y sensores perimetrales."),
            ("B2B/inmobiliario", "Agregar seguridad como diferencial.", "Analitica por edificio, barrio cerrado o proveedor."),
        ],
        [3.2, 5.5, 7.1],
    )

    doc.add_heading("4. Analisis FODA", level=1)
    add_table(
        doc,
        ["Fortalezas", "Oportunidades"],
        [
            ("Integra senales visuales, auditivas, sensoriales y contextuales.", "Crecimiento de seguridad conectada y hogares inteligentes."),
            ("Reduce falsas alarmas antes de saturar la CRA.", "Expansion a segmentos Premium y Agro/Rural."),
            ("Genera trazabilidad explicable para operadores y auditoria.", "Alianzas con telcos, inmobiliarias e integradores."),
            ("Alimenta BI gerencial con datos operativos.", "Diferenciacion frente a kits autoinstalables."),
        ],
        [8, 8],
    )
    add_table(
        doc,
        ["Debilidades", "Amenazas"],
        [
            ("Requiere inversion inicial, capacitacion e integracion tecnica.", "Ciberataques, filtraciones o uso indebido de datos sensibles."),
            ("Puede generar resistencia si se percibe como reemplazo humano.", "Competencia de Big Tech y soluciones economicas."),
            ("La interoperabilidad con multiples dispositivos eleva complejidad.", "Cambios regulatorios sobre privacidad y decisiones automatizadas."),
            ("Debe ser explicable para evitar decisiones opacas.", "Costos de hardware importado y contexto macroeconomico."),
        ],
        [8, 8],
    )
    add_p(doc, "Conclusion FODA: la propuesta es estrategicamente atractiva si se implementa de manera gradual, con supervision humana, explicabilidad, auditoria y controles de privacidad desde el diseno.")

    doc.add_heading("5. Organigrama, funciones y puestos", level=1)
    add_p(doc, "El area analizada es Operaciones y Central Receptora de Alarmas, porque alli Smart-Twin impacta directamente: recepcion de senales, clasificacion, validacion, escalamiento y registro del evento.")
    add_table(
        doc,
        ["Area/Puesto", "Funcion principal"],
        [
            ("Direccion de Operaciones", "Define estrategia operativa, presupuesto, SLA y calidad de respuesta."),
            ("Gerencia CRA", "Administra operadores, protocolos, turnos y escalamiento."),
            ("Coordinador de turno", "Supervisa eventos complejos y distribuye carga."),
            ("Operador CRA Nivel 1", "Valida eventos verdes/amarillos y contacta clientes."),
            ("Operador CRA Nivel 2", "Gestiona alertas rojas, eventos criticos y derivaciones."),
            ("Analista BI", "Construye KPIs, tableros y reportes gerenciales."),
            ("Especialista IA/Automatizacion", "Ajusta reglas, revisa sesgos y documenta modelos."),
            ("Ciberseguridad/Privacidad", "Controla accesos, cifrado, auditoria y cumplimiento."),
            ("Tecnico instalador", "Instala sensores, camaras, conectividad y soporte en campo."),
        ],
        [4.5, 11.5],
    )

    doc.add_heading("6. Procedimiento operativo propuesto", level=1)
    add_numbered(
        doc,
        [
            "Un dispositivo detecta una senal visual, auditiva, sensorial o de conectividad.",
            "El nodo Edge realiza una primera interpretacion conceptual para reducir latencia y volumen de datos.",
            "Smart-Twin construye una representacion digital del evento y calcula un Risk Score.",
            "El sistema clasifica la alarma: verde, amarilla o roja.",
            "Las verdes se registran como falsas alarmas filtradas; las amarillas pasan a revision humana; las rojas se priorizan.",
            "El evento alimenta el historial operativo y el tablero BI gerencial.",
        ],
    )

    doc.add_heading("7. Tres estrategias para lograr objetivos", level=1)
    add_table(
        doc,
        ["Estrategia", "Acciones", "Resultado esperado"],
        [
            ("Seguridad predictiva con IA y Edge", "Implementar scoring multimodal, reglas explicables, hand-off sensorial y procesamiento local conceptual.", "Menos falsas alarmas, menor tiempo de clasificacion y mayor trazabilidad."),
            ("Expansion Premium y Agro/Rural", "Crear planes con LoRaWAN, satelital, sensores perimetrales, camara termica y reportes avanzados.", "Nuevo ingreso, mayor ticket promedio y cobertura de zonas con baja conectividad."),
            ("Ecosistema B2B + BI", "Integrar proveedores, camaras compatibles, tableros gerenciales y analitica por zona/segmento.", "Decisiones gerenciales basadas en datos y diferenciacion competitiva."),
        ],
        [4, 7, 5],
    )

    doc.add_heading("8. Aplicacion tecnologica", level=1)
    add_table(
        doc,
        ["Tecnologia", "Aplicacion en Smart-Twin", "Justificacion"],
        [
            ("TIC/NTIC", "Sensores, paneles, app, software CRA, redes y tableros.", "Digitalizan procesos y mejoran coordinacion operativa."),
            ("BI", "KPIs de falsas alarmas, despachos, zonas, tiempos y ahorro.", "Permite decisiones gerenciales basadas en evidencia."),
            ("Cloud Computing", "Historicos, dashboards, data warehouse y escalabilidad.", "Absorbe picos de eventos y facilita analitica."),
            ("Virtualizacion/Kubernetes", "Contenedores y despliegues escalables conceptuales.", "Mejora resiliencia, mantenimiento y portabilidad."),
            ("IA", "Motor de Risk Scoring explicable; evolucion futura a ML supervisado.", "Prioriza eventos y reduce ruido operativo."),
            ("IoT", "Sensores de apertura, movimiento, shock, inhibicion y perimetro.", "Convierte senales fisicas en informacion accionable."),
            ("BYOD + MDM", "Uso controlado para comerciales y gerencia; restringido en CRA.", "Movilidad con gobierno de datos y seguridad."),
            ("RA/RV", "RA para instalacion tecnica; RV para capacitacion CRA.", "Reduce errores y mejora entrenamiento."),
            ("Zero Trust", "Autenticacion fuerte, roles, auditoria y cifrado.", "Protege datos sensibles y cumple buenas practicas."),
        ],
        [3.2, 7.3, 5.5],
    )
    doc.add_heading("8.1 Ciberseguridad, privacidad y Ley 25.326", level=2)
    add_p(doc, "Smart-Twin procesaria datos sensibles: horarios, ubicacion, imagenes, audio y patrones de presencia. Por eso, una implementacion real debe aplicar privacidad desde el diseno, minimizacion, consentimiento informado, cifrado, autenticacion multifactor, control de acceso por roles, auditoria y retencion limitada. El MVP no almacena imagenes reales, no usa reconocimiento facial y no envia alertas reales. La Ley 25.326 exige proteger datos personales y derechos como intimidad y honor.")

    doc.add_heading("9. Prototipo funcional / MVP", level=1)
    add_p(doc, "El MVP es una aplicacion web local sin backend real ni llamadas a APIs externas. Fue construido para mostrar la funcionalidad principal: simular eventos, ejecutar un analisis de riesgo, explicar la decision y alimentar indicadores gerenciales en tiempo real.")
    add_table(
        doc,
        ["Pantalla", "Funcion"],
        [
            ("Operador CRA", "Carga eventos por selectores, ejecuta el analisis, muestra score, estado, explicacion, accion recomendada, trazabilidad y payload simulado en alertas rojas."),
            ("Gerencia", "Muestra KPIs y graficos alimentados solo por las simulaciones de la sesion: eventos, falsas alarmas, despachos, revisiones, respuesta y ahorro."),
        ],
        [4, 12],
    )
    add_table(
        doc,
        ["Rango", "Estado", "Interpretacion"],
        [
            ("0 a 25", "Verde", "Falsa alarma filtrada o baja criticidad."),
            ("26 a 59", "Amarillo", "Revision humana requerida."),
            ("60 a 100", "Rojo", "Intrusion critica verificada o evento prioritario."),
        ],
        [3, 4, 9],
    )
    add_p(doc, "Casos demo: mascota + ladrido; persona desconocida + movimiento; arma + rotura de cristal; ZeroVision con camara obstruida; sensor rural + inhibicion de senal. El payload 911 es solo academico y no se envia a autoridades.")

    doc.add_heading("10. Viabilidad, roadmap y conclusion", level=1)
    add_table(
        doc,
        ["Dimension", "Evaluacion"],
        [
            ("Tecnica", "Viable con sensores IoT, Edge, Cloud, dashboards BI, contenedores y controles de seguridad. El reto esta en la integracion confiable."),
            ("Economica", "La inversion se justifica por ahorro operativo, menor saturacion CRA y nuevos planes Premium/Agro."),
            ("Organizacional", "Requiere capacitacion y gestion del cambio: Smart-Twin asiste al operador, no lo reemplaza."),
            ("Legal/etica", "Necesita consentimiento, minimizacion de datos, auditoria, explicabilidad y cumplimiento normativo."),
        ],
        [3.5, 12.5],
    )
    add_table(
        doc,
        ["Fase", "Alcance", "Resultado"],
        [
            ("1. MVP academico", "Demo local con datos simulados.", "Validacion conceptual para la materia."),
            ("2. Piloto interno", "Eventos historicos anonimizados y operadores seleccionados.", "Ajuste de reglas y medicion de impacto."),
            ("3. Piloto cliente Premium", "Prueba limitada con consentimiento y supervision humana.", "Validacion comercial y operativa."),
            ("4. Escalado Cloud", "Contenedores, Kubernetes, BI corporativo y seguridad avanzada.", "Operacion escalable y auditable."),
            ("5. Expansion Agro/B2B", "LoRaWAN, satelital, alianzas e integraciones.", "Nuevo segmento de negocio."),
        ],
        [3, 7, 6],
    )
    add_p(doc, "Conclusion: Verisure Smart-Twin es una propuesta de transformacion digital aplicable al sector de seguridad monitoreada. Integra tendencias tecnologicas con un problema operativo concreto: reducir falsas alarmas, priorizar eventos criticos y generar inteligencia gerencial. Su principal fortaleza es combinar eficiencia, trazabilidad y responsabilidad: decide antes, responde mejor y mantiene supervision humana en los casos sensibles.")

    doc.add_heading("Fuentes consultadas", level=1)
    add_table(
        doc,
        ["Ref.", "Fuente / contenido"],
        [
            ("F1", "Verisure Argentina - Central de monitoreo 24/7: https://www.verisure.com.ar/servicios/central-de-alarmas"),
            ("F2", "Verisure Argentina - ZeroVision: https://www.verisure.com.ar/productos/zerovision"),
            ("F3", "Argentina.gob.ar - Ley 25.326 de Proteccion de Datos Personales: https://www.argentina.gob.ar/normativa/nacional/ley-25326-64790/texto"),
            ("F4", "AWS - What is cloud computing: https://docs.aws.amazon.com/whitepapers/latest/aws-overview/what-is-cloud-computing.html"),
            ("F5", "Docker - What is a container: https://www.docker.com/resources/what-container"),
            ("F6", "LoRa Alliance - About LoRaWAN: https://lora-alliance.org/about-lorawan/"),
            ("F7", "ONVIF - Perfiles de interoperabilidad para seguridad IP: https://www.onvif.org"),
            ("F8", "Microsoft - Administracion de dispositivos / MDM: https://learn.microsoft.com"),
        ],
        [2, 14],
    )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(8.7)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    if OUT.exists():
        OUT.unlink()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
